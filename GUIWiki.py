# GUIWiki.py
import wikipedia #

#python to docx
from docx import Document
def Wiki (keyword,lang='th'):
    wikipedia.set_lang(lang)
    
    # summary สำหรับบทความที่สรุป
    data = wikipedia.summary(keyword)

    # page + comtent บทความทั้งหน้า
    data2 = wikipedia.page(keyword)
    data2 = data2.content

    

    doc = Document() #สร้างไฟล์ word ใน python
    doc.add_heading(keyword,0)

    #messagebox.showinfo('Search complete',data2)
    
    doc.add_paragraph(data2)
    doc.save(keyword + '.docx')
    
    print('สร้างไฟล์สำเร็จ')
    

#เปลี่ยนเป็นภาษาไทย
wikipedia.set_lang('th')

from tkinter import * # import * เป็นการ import คำสั่งทั้งหมดของฟังก์ชัน Main
from tkinter import ttk
from tkinter import messagebox

guis = Tk()
guis.title('โปรแกรม wiki')
guis.geometry('400x300') # Size ของหน้าต่างโปรแกรม

#config
FONT1 = ('Angsana New',15) # set font

#คำอธิบาย

L = ttk.Label(guis, text = 'ค้นหาบทความ',font = FONT1) #Label ข้อความ
L.pack()

#ช่องค้นหาข้อมูล
v_search = StringVar() # กล่องเก็บคำที่ต้องการค้นหา
E1 = ttk.Entry(guis,textvariable = v_search,font = FONT1,width=40)
E1.pack(pady=10)


#ปุ่มค้นหา

def Search():
    keyword = v_search.get() # .get() ใช้ดึงข้อมูลจากตัวแปรประเภท StringVar() เท่านั้น
    try:
        #ลองค้นหาดูว่าได้ผลลัพธ์หรือไม่ หากได้ให้ผ่านไป
        language = v_radio.get()
        Wiki(keyword,language)
        messagebox.showinfo('Search complete','ค้นหาข้อความสำเร็จ บันทึกเรียบร้อยแล้ว')
    except:
        #หากรับคำสั่งแล้วมีปัญหา แสดงข้อความแจ้งเตือน
        messagebox.showwarning('Keyword Error','กรุณากรอกคำค้นหาใหม่')

        
    #print(wikipedia.search(keyword))
    #result = wikipedia.summary(keyword)
    #print(result)
    
B1 = ttk.Button(guis,text = 'Search',command = Search)
B1.pack(ipadx=20,ipady=10)

#เลือกภาษา
F1 = Frame(guis)
F1.pack(pady=10)

v_radio = StringVar()

RB1 = ttk.Radiobutton(F1,text = 'ภาษาไทย', variable = v_radio,value = 'th')
RB2 = ttk.Radiobutton(F1,text = 'English', variable = v_radio,value = 'en')
RB3 = ttk.Radiobutton(F1,text = '日本', variable = v_radio,value = 'jp')

RB1.grid(row=0,column=0)
RB2.grid(row=0,column=1)
RB3.grid(row=0,column=2)


guis.mainloop()
